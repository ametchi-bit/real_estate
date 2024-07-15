import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml import OxmlElement
from datetime import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
import plotly.express as px
from plotly.offline import plot
import plotly.io as pio
import tempfile
from docx2pdf import convert
from io import BytesIO
import io
import pythoncom
import win32com.client
from util import summary_desc_symphonia, explain_summary_desc, symphonia_data


def export_to_docx_pdf(
    title,
    subtitle,
    intro_title,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    col_data_desc_col,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    desc_stat_content,
    # data_distribution,
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
):

    # Function to set column widths in a table
    def set_column_widths(table, widths):
        for i, width in enumerate(widths):
            if isinstance(width, pd.DataFrame):
                # Calculate the maximum width based on the content of the DataFrame
                max_width = Inches(width.astype(str).applymap(len).max().max())
                table.cell(0, i).width = max_width
            else:
                # Assume width is a valid length unit (e.g., Inches(1))
                table.cell(0, i).width = width

    # Function to create a custom table style
    def create_custom_table_style(doc, style_name, grid_color_rgb=(0, 0, 0)):
        style_element = OxmlElement("w:style")
        style_element.set(qn("w:type"), "table")
        style_element.set(qn("w:styleId"), style_name)

        name_element = OxmlElement("w:name")
        name_element.set(qn("w:val"), style_name)
        style_element.append(name_element)

        based_on_element = OxmlElement("w:basedOn")
        based_on_element.set(qn("w:val"), "TableGrid")
        style_element.append(based_on_element)

        paragraph_properties_element = OxmlElement("w:pPr")
        spacing_element = OxmlElement("w:spacing")
        spacing_element.set(qn("w:before"), "0")
        spacing_element.set(qn("w:after"), "0")
        paragraph_properties_element.append(spacing_element)
        style_element.append(paragraph_properties_element)

        table_properties_element = OxmlElement("w:tblPr")
        table_style_element = OxmlElement("w:tblStyle")
        table_style_element.set(qn("w:val"), style_name)
        table_properties_element.append(table_style_element)

        grid_element = OxmlElement("w:tblBorders")
        grid_element.set(qn("w:val"), "grid")
        grid_element.set(qn("w:color"), f"auto")
        grid_element.set(qn("w:sz"), "4")
        grid_element.set(qn("w:space"), "0")
        table_properties_element.append(grid_element)

        style_element.append(table_properties_element)

        color_element = OxmlElement("w:shd")
        color_element.set(
            qn("w:fill"),
            f"{grid_color_rgb[0]:02X}{grid_color_rgb[1]:02X}{grid_color_rgb[2]:02X}",
        )
        style_element.append(color_element)

        doc.styles._element.append(style_element)
        return style_name  # Fix here

    # Create a new Word document
    doc = Document("report_template.docx")

    # Iterate through styles and set font name
    for style in doc.styles:
        if "Normal" in style.name:
            style.font.name = "Times New Roman"

    # Create custom table styles
    table_style_name = create_custom_table_style(doc, "TableGrid4Accent5")

    # Set top margin for the header
    section = doc.sections[0]
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    # page title
    page_title = doc.add_heading(0)
    page_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_title_run = page_title.add_run(title)
    page_title_run.bold = True

    # page subtitle
    page_subtitle = doc.add_heading(level=1)
    page_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_subtitle_run = page_subtitle.add_run(subtitle)
    page_subtitle_run.bold = True

    # Introduction title
    introduction_title = doc.add_heading(level=1)
    introduction_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    introduction_title_run = introduction_title.add_run(intro_title)
    introduction_title_run.bold = True

    # Objectif title
    objectif_title = doc.add_heading(level=2)
    objectif_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    objectif_title_run = objectif_title.add_run(title_objectif)
    objectif_title_run.bold = True

    # Objectif text
    objectif_text = doc.add_paragraph()
    objectif_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    objectif_text_run = objectif_text.add_run(intro_objectif)
    objectif_text_run.font.size = Pt(14)

    # Scope title
    scope_title = doc.add_heading(level=2)
    scope_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    scope_title_run = scope_title.add_run(title_scope)
    scope_title_run.bold = True

    # Scope text
    scope_text = doc.add_paragraph()
    scope_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    scope_text_run = scope_text.add_run(text_scope)
    scope_text_run.bold = True

    # Description title
    description_title = doc.add_heading(level=2)
    description_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description_title_run = description_title.add_run(title_data_desc)
    description_title_run.bold = True

    # Description subtitle
    description_subtitle = doc.add_heading(level=3)
    description_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description_subtitle_run = description_subtitle.add_run(sub_title_data_text)
    description_subtitle_run.bold = True

    # Description sub content 1
    description_sub_content = doc.add_paragraph()
    description_sub_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    description_sub_content_run = description_sub_content.add_run(data_overview_text)
    description_sub_content_run.font.size = Pt(14)

    # Description subtitle 2
    description_subtitle2 = doc.add_heading(level=3)
    description_subtitle2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description_subtitle2_run = description_subtitle2.add_run(sub_title_data_desc_col)
    description_subtitle2_run.bold = True

    # Description sub content 2
    # description_sub_content_2 = doc.add_paragraph()
    # description_sub_content_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # description_sub_content_2_run = description_sub_content_2.add_run(col_data_desc_col)
    # description_sub_content_2_run.font.size = Pt(14)

    column_descriptions = {
        "Horodatage": "Horodatage de l'enregistrement de données",
        "Caméra": "Identifiant de la caméra",
        "Scénario": "Description du scénario",
        "Catégorie": "Catégorie de l'objet détecté",
        "Type d'objet": "Type d'objet détecté",
        "Nombre": "Nombre d'objets détectés",
    }

    def get_column_descriptions(symphonia_data):
        columns = symphonia_data.columns
        descriptions = {}
        for col in columns:
            description = column_descriptions.get(col, "No description available")
            descriptions[col] = description
        return descriptions

    def display_column_descriptions(symphonia_data):
        descriptions = get_column_descriptions(symphonia_data)
        for col, desc in descriptions.items():
            print(f"**{col}**: {desc}")

    # Add column descriptions
    description_paragraph = doc.add_paragraph()
    description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Display column descriptions
    descriptions = get_column_descriptions(symphonia_data)
    for col, desc in descriptions.items():
        description_paragraph.add_run(f"{col}: {desc}\n").font.size = Pt(14)

    # Data cleaning title
    data_cleanind_prep_title = doc.add_heading(level=2)
    data_cleanind_prep_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_cleanind_prep_title_run = data_cleanind_prep_title.add_run(data_cleaning_title)
    data_cleanind_prep_title_run.bold = True

    # Data cleaning subtitle missing value
    data_cleanind_prep_subtitle = doc.add_heading(level=3)
    data_cleanind_prep_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_cleanind_prep_subtitle_run = data_cleanind_prep_subtitle.add_run(
        missing_val_title
    )
    data_cleanind_prep_subtitle_run.bold = True

    # Data cleaning content missing value
    data_missing_text = doc.add_paragraph()
    data_missing_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_missing_text_run = data_missing_text.add_run(missing_val_content)
    data_missing_text_run.font.size = Pt(14)

    # Data cleaning subtitle transformation
    data_transformation = doc.add_heading(level=3)
    data_transformation.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_transformation_run = data_transformation.add_run(transformation_data_title)
    data_transformation_run.bold = True

    # Data cleaning content transformation
    data_transformation_text = doc.add_paragraph()
    data_transformation_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_transformation_text_run = data_transformation_text.add_run(
        transformation_data_content
    )
    data_transformation_text_run.font.size = Pt(14)

    # Descriptive statistics title
    descriptive_stat_title = doc.add_heading(level=2)
    descriptive_stat_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    descriptive_stat_title_run = descriptive_stat_title.add_run(desc_stat_title)
    descriptive_stat_title_run.bold = True

    # Summary description

    # Add summary statistics table
    summary_data = summary_desc_symphonia(symphonia_data)
    explanations = explain_summary_desc(summary_data)

    # Add table with summary statistics

    table = doc.add_table(
        rows=summary_data.shape[0] + 1,
        cols=summary_data.shape[1],
        style=table_style_name,
    )
    set_column_widths(
        table,
        [
            Inches(1),
            Inches(2),
            Inches(3),
            Inches(1),
            Inches(3),
            Inches(1),
            Inches(1),
            Inches(1),
            Inches(1),
        ],
    )
    # Add column headers to the table
    for col, header in enumerate(summary_data.columns):
        cell = table.cell(0, col)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add data to the table
    for row in range(summary_data.shape[0]):
        for col in range(summary_data.shape[1]):
            cell = table.cell(row + 1, col)
            cell.text = str(summary_data.iloc[row, col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            if col == 0:  # First column
                cell.paragraphs[0].paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            else:
                cell.paragraphs[0].paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )
    # Add explanation_sum_desc

    explanation_sum_desc = doc.add_paragraph()
    explanation_sum_desc.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Display column descriptions
    explanations = explain_summary_desc(summary_data)
    for col, desc in explanations.items():
        explanation_sum_desc.add_run(f"{col}: {desc}\n").font.size = Pt(14)

    # Distribution dataframe
    # Distribution plot
    distribution_fig = px.histogram(
        symphonia_data,
        x="Type d'objet",
        y="Nombre",
        title="Distribution des données",
    )

    # Convert Plotly figure to image bytes
    image_stream = BytesIO()
    pio.write_image(
        distribution_fig, image_stream, format="png", width=955, height=525, scale=2
    )
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(8), height=Inches(8))

    # Conclusion
    conclusion_main_title = doc.add_heading(level=2)
    conclusion_main_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_main_title_run = conclusion_main_title.add_run(conclusion_title)
    conclusion_main_title_run.bold = True

    # Conclusion resume
    conclusion_resume = doc.add_heading(level=3)
    conclusion_resume.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_resume_run = conclusion_resume.add_run(conclusion_subtitl_r)
    conclusion_resume_run.bold = True

    # Conclusion resume content
    conclusion_resume_content = doc.add_paragraph()
    conclusion_resume_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conclusion_resume_content_run = conclusion_resume_content.add_run(resume_content)
    conclusion_resume_content_run.font.size = Pt(14)

    # Conclusion recommendation titre
    conclusion_recomm = doc.add_heading(level=3)
    conclusion_recomm.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conclusion_recomm_run = conclusion_recomm.add_run(recomm_title)
    conclusion_recomm_run.bold = True

    # Conclusion resume content
    conclusion_recomm_content = doc.add_paragraph()
    conclusion_recomm_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conclusion_recomm_content_run = conclusion_recomm_content.add_run(recomm_content)
    conclusion_recomm_content_run.font.size = Pt(14)

    # Auto-save function using raport title and timestamp
    timestamp = datetime.now().strftime("%d-%m-%Y %H-%M-%S")
    auto_save_filename = f"{title}-{timestamp}.docx"
    pdf_filename = f"{title}-{timestamp}.pdf"

    # Save the document
    doc.save(auto_save_filename)

    def convert_docx_to_pdf(docx_file, pdf_file):
        pythoncom.CoInitialize()
        try:
            convert(docx_file, pdf_file)
        finally:
            pythoncom.CoUninitialize()

    # Convert the DOCX to PDF
    convert_docx_to_pdf(auto_save_filename, pdf_filename)

    # return doc as buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_report(
    title,
    subtitle,
    intro_title,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    col_data_desc_col,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    desc_stat_content,
    # data_distribution,
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
):
    pdf_data = export_to_docx_pdf(
        title,
        subtitle,
        intro_title,
        title_objectif,
        intro_objectif,
        title_scope,
        text_scope,
        title_data_desc,
        sub_title_data_text,
        data_overview_text,
        sub_title_data_desc_col,
        col_data_desc_col,
        data_cleaning_title,
        missing_val_title,
        missing_val_content,
        transformation_data_title,
        transformation_data_content,
        desc_stat_title,
        desc_stat_content,
        # data_distribution,
        conclusion_title,
        conclusion_subtitl_r,
        resume_content,
        recomm_title,
        recomm_content,
    )
    return pdf_data
