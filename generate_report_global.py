import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml import OxmlElement
from datetime import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
import plotly.express as px
from plotly.offline import plot
import plotly.io as pio
import tempfile
from docx2pdf import convert
from io import BytesIO
import plotly.graph_objects as go
import io
import pythoncom
import win32com.client
import uuid
from util import (analyse_comparative_par_camera, analyse_correlation, analyze_anomalies, diagrammes_de_dispersion, ecart_type_intercameras, segmentation_grouping, summary_desc_symphonia, explain_summary_desc, symphonia_data, repartition_viz, plot_comparaison_type_objet_real_estate, plot_weekly_daily_patterns_real_estate)


def export_to_docx_pdf(
    title,
    subtitle,
    intro_title,
    intro_brief,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    ms_subtitle_def_typob,
    text_def_to,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    desc_stat_subtite,
    desc_stat_rgd_content,
    desc_stat_content,
    ms_distribution_title,
    ms_tendances_temporelles,
    ms_variabilite_cameras,
    ms_ecart_type_intercam,
    data_viz,
    ms_repartition_camera,
    ms_corr_analysis,
    ms_anomali_detection,
    ms_seg_group,
    
    # interpretation et synthese
    ms_interpretation_synthese_title,
    ms_summary_subtitle,
    ms_summary_text,
    ms_summary_result, 
            
    #interpretation
    ms_interpretation_subtitle,
    ms_interpretation_text,
    ms_interpretation_result,
    
    # lien avec l'objet du rapport
    ms_lien_rapport_subtitle,
    ms_lien_rapport,
            
    #conclusion preliminaire
    ms_conclusison_pre_subtitle,
    ms_conclusion_pre,
    
    #conclusion
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
    ms_point_final,
):

    # Function to set column widths in a table
    def set_column_widths(table, widths):
        for i, width in enumerate(widths):
            if isinstance(width, pd.DataFrame):
                # Calculate the maximum width based on the content of the DataFrame
                max_width = Inches(width.astype(str).applymap(len).max().max())
                table.cell(0, i).width = max_width
            else:
                # Assume width is a valid length unit (e.g., Inches(1))
                table.cell(0, i).width = width

    # Function to create a custom table style
    def create_custom_table_style(doc, style_name, grid_color_rgb=(0, 0, 0)):
        style_element = OxmlElement("w:style")
        style_element.set(qn("w:type"), "table")
        style_element.set(qn("w:styleId"), style_name)

        name_element = OxmlElement("w:name")
        name_element.set(qn("w:val"), style_name)
        style_element.append(name_element)

        based_on_element = OxmlElement("w:basedOn")
        based_on_element.set(qn("w:val"), "TableGrid")
        style_element.append(based_on_element)

        paragraph_properties_element = OxmlElement("w:pPr")
        spacing_element = OxmlElement("w:spacing")
        spacing_element.set(qn("w:before"), "0")
        spacing_element.set(qn("w:after"), "0")
        paragraph_properties_element.append(spacing_element)
        style_element.append(paragraph_properties_element)

        table_properties_element = OxmlElement("w:tblPr")
        table_style_element = OxmlElement("w:tblStyle")
        table_style_element.set(qn("w:val"), style_name)
        table_properties_element.append(table_style_element)

        grid_element = OxmlElement("w:tblBorders")
        grid_element.set(qn("w:val"), "grid")
        grid_element.set(qn("w:color"), f"auto")
        grid_element.set(qn("w:sz"), "4")
        grid_element.set(qn("w:space"), "0")
        table_properties_element.append(grid_element)

        style_element.append(table_properties_element)

        color_element = OxmlElement("w:shd")
        color_element.set(
            qn("w:fill"),
            f"{grid_color_rgb[0]:02X}{grid_color_rgb[1]:02X}{grid_color_rgb[2]:02X}",
        )
        style_element.append(color_element)

        doc.styles._element.append(style_element)
        return style_name  # Fix here

    # Create a new Word document
    doc = Document("report_template.docx")

    # Iterate through styles and set font name
    for style in doc.styles:
        if "Normal" in style.name:
            style.font.name = "Times New Roman"

    # Create custom table styles
    table_style_name = create_custom_table_style(doc, "TableGrid4Accent5")

    # Set top margin for the header
    section = doc.sections[0]
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    

    # page title
    page_title = doc.add_heading(0)
    page_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_title_run = page_title.add_run(title)
    page_title_run.bold = True
    

    # page subtitle
    page_subtitle = doc.add_heading(level=1)
    page_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_subtitle_run = page_subtitle.add_run(subtitle)
    page_subtitle_run.bold = True
    doc.add_page_break()
    
    # Introduction title
    introduction_title = doc.add_heading(level=1)
    introduction_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    introduction_title_run = introduction_title.add_run(intro_title)
    introduction_title_run.bold = True

    #intro_brief text
    intro_brief_text = doc.add_paragraph()
    intro_brief_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_brief_text_run =intro_brief_text.add_run(intro_brief)
    intro_brief_text_run.font.size = Pt(14)

    # Objectif title
    objectif_title = doc.add_heading(level=2)
    objectif_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    objectif_title_run = objectif_title.add_run(title_objectif)
    objectif_title_run.bold = True

    # Objectif text
    objectif_text = doc.add_paragraph()
    objectif_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    objectif_text_run = objectif_text.add_run(intro_objectif)
    objectif_text_run.font.size = Pt(14)

    # Scope title
    scope_title = doc.add_heading(level=2)
    scope_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    scope_title_run = scope_title.add_run(title_scope)
    scope_title_run.bold = True

    # Scope text
    scope_text = doc.add_paragraph()
    scope_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    scope_text_run = scope_text.add_run(text_scope)
    scope_text_run.font.size = Pt(14)

    # Description title
    description_title = doc.add_heading(level=2)
    description_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    description_title_run = description_title.add_run(title_data_desc)
    description_title_run.bold = True

    # Description subtitle
    description_subtitle = doc.add_heading(level=3)
    description_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    description_subtitle_run = description_subtitle.add_run(sub_title_data_text)
    description_subtitle_run.bold = True

    # Adding the description sub content
    description_sub_content = doc.add_paragraph()
    description_sub_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    description_sub_content_run = description_sub_content.add_run(data_overview_text)
    description_sub_content_run.font.size = Pt(14)

    # Description subtitle 2
    description_subtitle2 = doc.add_heading(level=3)
    description_subtitle2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    description_subtitle2_run = description_subtitle2.add_run(sub_title_data_desc_col)
    description_subtitle2_run.bold = True
    
    # Define a custom bullet list style
    styles = doc.styles
    style = styles.add_style('MyListBullet', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['List Paragraph']
    p = style.paragraph_format
    p.left_indent = Pt(36)
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    
    # Bulleted list content
    bullet_points = [
        "Horodatage : Cette colonne indique la date et l'heure précises de chaque événement enregistré. Elle est essentielle pour analyser les tendances temporelles et les schémas de comportement.",
        "Caméra : Identifie la caméra spécifique ayant capturé l'événement. Cela permet de localiser géographiquement les événements et de comprendre les zones de la propriété où l'activité est la plus dense.",
        "Scénario : Décrit le contexte ou la situation dans laquelle l'événement s'est produit, aidant ainsi à catégoriser les différentes situations de surveillance.",
        "Catégorie : Classe les événements en différentes catégories, facilitant ainsi une analyse segmentée des données.",
        "Type d'objet : Indique le type d'objet détecté, par exemple, personne, moto, véhicule léger, etc. Cette colonne est cruciale pour comprendre la distribution des différents objets sur la propriété.",
        "Nombre : Quantifie le nombre d'objets détectés pour chaque enregistrement, permettant une analyse quantitative des données."
    ]

    # Adding the bullet points to the document
    for point in bullet_points:
        p = doc.add_paragraph(style='MyListBullet')
        run = p.add_run(point)
        run.font.size = Pt(14)

    doc.add_paragraph()
    

    # Description subtitle 3
    description_subtitle3 = doc.add_heading(level=3)
    description_subtitle3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    description_subtitle3_run = description_subtitle3.add_run(ms_subtitle_def_typob)
    description_subtitle3_run.bold = True
    
    # Define a custom bullet list style
    styles2 = doc.styles
    style = styles2.add_style('ListBullet', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles2['List Paragraph']
    p2 = style.paragraph_format
    p2.left_indent = Pt(36)
    p2.space_before = Pt(0)
    p2.space_after = Pt(0)
    
    # Bulleted list content
    bullet_points = [
    "Bus : Véhiculent des passagers en grands nombres et nécessitent une surveillance particulière en raison de leur taille et de leur impact potentiel sur la circulation."
    " Moto : Inclut toutes les formes de motocyclettes, souvent plus rapides et plus difficiles à détecter que les véhicules plus grands."
    " Personne : Comprend tous les piétons, offrant des insights sur les mouvements des résidents, visiteurs et personnels."
    "Véhicule intermédiaire : Réfère aux véhicules de taille moyenne, tels que les camionnettes, qui jouent un rôle clé dans la logistique et le transport."
    "Véhicule léger : Inclut les voitures personnelles et autres véhicules de petite taille, couramment utilisés par les résidents et visiteurs."
    "Vélo : Inclut toutes les formes de bicyclettes, de plus en plus courantes dans les zones résidentielles et nécessitant une surveillance pour la sécurité."
    "Poid lourd : Véhicule de grande taille destiné au transport de marchandises, caractérisé par une capacité de charge élevée."

    ]

    # Adding the bullet points to the document
    for point in bullet_points:
        p2 = doc.add_paragraph(style='ListBullet')
        run = p2.add_run(point)
        run.font.size = Pt(14)

    doc.add_paragraph()

    # Data cleaning title
    data_cleanind_prep_title = doc.add_heading(level=2)
    data_cleanind_prep_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_cleanind_prep_title_run = data_cleanind_prep_title.add_run(data_cleaning_title)
    data_cleanind_prep_title_run.bold = True

    # Data cleaning subtitle missing value
    data_cleanind_prep_subtitle = doc.add_heading(level=3)
    data_cleanind_prep_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_cleanind_prep_subtitle_run = data_cleanind_prep_subtitle.add_run(
        missing_val_title
    )
    data_cleanind_prep_subtitle_run.bold = True

    # Data cleaning content missing value
    data_missing_text = doc.add_paragraph()
    data_missing_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_missing_text_run = data_missing_text.add_run(missing_val_content)
    data_missing_text_run.font.size = Pt(14)

    # Data cleaning subtitle transformation
    data_transformation = doc.add_heading(level=3)
    data_transformation.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_transformation_run = data_transformation.add_run(transformation_data_title)
    data_transformation_run.bold = True

    # Data cleaning content transformation
    data_transformation_text = doc.add_paragraph()
    data_transformation_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    data_transformation_text_run = data_transformation_text.add_run(
        transformation_data_content
    )
    data_transformation_text_run.font.size = Pt(14)

    # Descriptive statistics title
    descriptive_stat_title = doc.add_heading(level=2)
    descriptive_stat_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    descriptive_stat_title_run = descriptive_stat_title.add_run(desc_stat_title)
    descriptive_stat_title_run.bold = True

    # Descriptive statistics subtitle resume general donees
    descriptive_stat_subtitle = doc.add_heading(level=2)
    descriptive_stat_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    descriptive_stat_subtitle_run = descriptive_stat_subtitle.add_run(desc_stat_subtite)
    descriptive_stat_subtitle_run.bold = True
    
    # Descriptive statistics subtitle resume general donees contenu
    descriptive_stat_sub_content = doc.add_paragraph()
    descriptive_stat_sub_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    descriptive_stat_sub_content_run = descriptive_stat_sub_content.add_run(desc_stat_rgd_content)
    descriptive_stat_sub_content_run.font.size = Pt(14)

    
    
    # Summary description

    # Add summary statistics table
    summary_data = summary_desc_symphonia(symphonia_data)

    # Add table with summary statistics

    table = doc.add_table(
        rows=summary_data.shape[0] + 1,
        cols=summary_data.shape[1],
        style=table_style_name,
    )
    set_column_widths(
        table,
        [
            Inches(1),
            Inches(2),
            Inches(3),
            Inches(1),
            Inches(3),
            Inches(1),
            Inches(1),
            Inches(1),
            Inches(1),
        ],
    )
    # Add column headers to the table
    for col, header in enumerate(summary_data.columns):
        cell = table.cell(0, col)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add data to the table
    for row in range(summary_data.shape[0]):
        for col in range(summary_data.shape[1]):
            cell = table.cell(row + 1, col)
            cell.text = str(summary_data.iloc[row, col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            if col == 0:  # First column
                cell.paragraphs[0].paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            else:
                cell.paragraphs[0].paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )
    doc.add_paragraph()
    # Add explanation_sum_desc


    # Data distribution title
    distribution_title = doc.add_heading(level=3)
    distribution_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    distribution_title_run = distribution_title.add_run(ms_distribution_title)
    distribution_title_run.bold = True
    

    # Distribution plot
    distribution_fig = px.histogram(
        symphonia_data,
        x="Type d'objet",
        y="Nombre",
        title="Distribution des données",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )

    # Convert Plotly figure to image bytes
    image_stream = BytesIO()
    pio.write_image(
        distribution_fig, image_stream, format="png", width=955, height=525, scale=2
    )
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(6.31), height=Inches(4.19))
    
    doc.add_page_break()

    # tendance temporel title
    tendances_temporel_title = doc.add_heading(level=3)
    tendances_temporel_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    tendances_temporel_title_run = tendances_temporel_title.add_run(ms_tendances_temporelles)
    tendances_temporel_title_run.bold = True
    
    # Trends 
    mwd_fig, image_stream_mwd = plot_weekly_daily_patterns_real_estate(symphonia_data, "Horodatage", "Nombre")
    doc.add_picture(image_stream_mwd, width=Inches(7.62), height=Inches(4.50))
    
    
    # variabilite et dispersion
    ana_compa_cam_fig, image_stream_compa_fig = analyse_comparative_par_camera(symphonia_data)

    analyse_comparative = doc.add_heading(level=2)
    analyse_comparative.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    analyse_comparative_run = analyse_comparative.add_run(data_viz)
    analyse_comparative_run.bold = True
    
    # Data distribution title
    variabilite_camera = doc.add_heading(level=3)
    variabilite_camera.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    variabilite_camera_run = variabilite_camera.add_run(ms_variabilite_cameras)
    variabilite_camera_run.bold = True
    
    # Analyse comparative du nombre moyen d'objet detecte par camera
    doc.add_picture(image_stream_compa_fig, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    # comparaison  des type d'objet
    comparaison_fig, image_stream_comparaison = plot_comparaison_type_objet_real_estate(symphonia_data)
    doc.add_picture(image_stream_comparaison, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    # Data distribution title
    ecartype_title = doc.add_heading(level=3)
    ecartype_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ecartype_title_run = ecartype_title.add_run(ms_ecart_type_intercam)
    ecartype_title_run.bold = True
    
    # ecart type d'objet
    ecart_type_fig, image_stream_ecart_type = ecart_type_intercameras(symphonia_data)
    doc.add_picture(image_stream_ecart_type, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    
     # Data distribution title
    repartition_title = doc.add_heading(level=3)
    repartition_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    repartition_title_run = repartition_title.add_run( ms_repartition_camera)
    repartition_title_run.bold = True
    
    # repartition des cameras
    rep_cam_fig, rep_scen_fig, fig_categorie,image_stream_cam,image_stream_cat, image_stream_scen = repartition_viz(symphonia_data)
    doc.add_picture(image_stream_cam, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    # repartition des detection d'objet par camera
    dia_disp_fig, image_stream_diagram_dispersion = diagrammes_de_dispersion(symphonia_data)
    doc.add_picture(image_stream_diagram_dispersion, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()

    
    # Data distribution title
    corr_analysis_title = doc.add_heading(level=3)
    corr_analysis_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_analysis_title_run = corr_analysis_title.add_run( ms_corr_analysis)
    corr_analysis_title_run.bold = True
    
    # analyse de correlation 
    analyse_corr_fig, image_stream_analyse_corr_fig = analyse_correlation(symphonia_data)
    doc.add_picture(image_stream_analyse_corr_fig, width=Inches(7.42), height=Inches(4.50))
    doc.add_page_break()
    
    # Data distribution title
    anomali_detection_title = doc.add_heading(level=3)
    anomali_detection_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    anomali_detection_title_run = anomali_detection_title.add_run( ms_anomali_detection)
    anomali_detection_title_run.bold = True
    
    #anomalie detection
    anomaly_plot, anomaly_summary, image_stream_anomaly_plot = analyze_anomalies(symphonia_data, threshold=0.01)
    doc.add_picture(image_stream_anomaly_plot, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    # Data distribution title
    segmentation_title = doc.add_heading(level=3)
    segmentation_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    segmentation_title_run = segmentation_title.add_run( ms_seg_group)
    segmentation_title_run.bold = True
    
    # segmentation 
    seg_group_fig, clusters, image_stream_seg_group_fig = segmentation_grouping(symphonia_data, n_clusters=5)
    doc.add_picture(image_stream_seg_group_fig, width=Inches(7.62), height=Inches(4.50))
    doc.add_page_break()
    
    
    # interpretation et synthese
    synthese_interpretation_title = doc.add_heading(level=2)
    synthese_interpretation_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    synthese_interpretation_title_run = synthese_interpretation_title.add_run(ms_interpretation_synthese_title)
    synthese_interpretation_title_run.bold = True
    
    #summary subtitle
    summary_subtile = doc.add_heading(level=3)
    summary_subtile.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    summary_subtile_run = summary_subtile.add_run(ms_summary_subtitle)
    summary_subtile_run.bold = True
    
    #summary
    summary_content = doc.add_paragraph()
    summary_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    summary_content_run = summary_content.add_run( ms_summary_text)
    summary_content_run.font.size = Pt(14)
   
    doc.add_paragraph()
   
   #result
    summary_result_content = doc.add_paragraph()
    summary_result_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    summary_result_content_run = summary_result_content.add_run( ms_summary_result)
    summary_result_content_run.font.size = Pt(14) 
    
    doc.add_paragraph()
    
    #interpretation subtitle
    interpretation_subtitle = doc.add_heading(level=3)
    interpretation_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    interpretation_subtitle_run = interpretation_subtitle.add_run(ms_interpretation_subtitle)
    interpretation_subtitle_run.bold = True
    
    doc.add_paragraph()
    
    #interpretation
    interpretation_content = doc.add_paragraph()
    interpretation_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    interpretation_content_run = interpretation_content.add_run( ms_interpretation_text)
    interpretation_content_run.font.size = Pt(14) 
    
    doc.add_paragraph()
    
    #result
    interpretatin_result_content = doc.add_paragraph()
    interpretatin_result_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    interpretatin_result_content_run = interpretatin_result_content.add_run( ms_interpretation_result)
    interpretatin_result_content_run.font.size = Pt(14) 
    
    doc.add_paragraph()
    
    # lien avec rapport subtitle
    lienRapport_subtitle = doc.add_heading(level=3)
    lienRapport_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lienRapport_subtitle_run = lienRapport_subtitle.add_run(ms_lien_rapport_subtitle)
    lienRapport_subtitle_run.bold = True
    
    doc.add_paragraph()
    
    # lien avec l'objet du rapport
    lien_rapport_content = doc.add_paragraph()
    lien_rapport_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lien_rapport_content_run = lien_rapport_content.add_run( ms_lien_rapport)
    lien_rapport_content_run.font.size = Pt(14) 
    
    doc.add_paragraph()
    
    # lien avec rapport subtitle
    conclusion_pre_subtitle = doc.add_heading(level=3)
    conclusion_pre_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_pre_subtitle_run = conclusion_pre_subtitle.add_run(ms_conclusison_pre_subtitle)
    conclusion_pre_subtitle_run.bold = True
    
    doc.add_paragraph()
    
    
    #conclusion preliminaire
    conclusion_pre_content = doc.add_paragraph()
    conclusion_pre_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_pre_content_run = conclusion_pre_content.add_run( ms_conclusion_pre)
    conclusion_pre_content_run.font.size = Pt(14) 
    
    doc.add_paragraph()
    
    # Conclusion
    conclusion_main_title = doc.add_heading(level=2)
    conclusion_main_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_main_title_run = conclusion_main_title.add_run(conclusion_title)
    conclusion_main_title_run.bold = True

    doc.add_paragraph()
    
    # Conclusion resume
    conclusion_resume = doc.add_heading(level=3)
    conclusion_resume.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_resume_run = conclusion_resume.add_run(conclusion_subtitl_r)
    conclusion_resume_run.bold = True

    doc.add_paragraph()
    
    # Conclusion resume content
    conclusion_resume_content = doc.add_paragraph()
    conclusion_resume_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_resume_content_run = conclusion_resume_content.add_run(resume_content)
    conclusion_resume_content_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Conclusion recommendation titre
    conclusion_recomm = doc.add_heading(level=3)
    conclusion_recomm.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_recomm_run = conclusion_recomm.add_run(recomm_title)
    conclusion_recomm_run.bold = True
    
    doc.add_paragraph()
    
    # Conclusion resume content
    conclusion_recomm_content = doc.add_paragraph()
    conclusion_recomm_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_recomm_content_run = conclusion_recomm_content.add_run(recomm_content)
    conclusion_recomm_content_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # poin final de la conclusion
    point_final_content = doc.add_paragraph()
    point_final_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    point_final_content_run = point_final_content.add_run(ms_point_final)
    point_final_content_run.font.size = Pt(14)
    

    # Auto-save function using raport title and timestamp
    timestamp = datetime.now().strftime("%d-%m-%Y %H-%M-%S")
    unique_id = str(uuid.uuid4())[:8]
    auto_save_filename = f"{title}-{timestamp}.docx"
    pdf_filename = f"{title}-{timestamp}-{unique_id}.pdf"

    # Save the document as DOCX
    doc.save(auto_save_filename)

    def convert_docx_to_pdf(docx_file, pdf_file):
        pythoncom.CoInitialize()
        try:
            convert(docx_file, pdf_file)
        finally:
            pythoncom.CoUninitialize()

     # Convert the DOCX to PDF
    convert_docx_to_pdf(auto_save_filename, pdf_filename)

    # Read the PDF file and return its content
    with open(pdf_filename, "rb") as pdf_file:
        pdf_content = pdf_file.read()

    return pdf_content


def generate_report(
    title,
    subtitle,
    intro_title,
    intro_brief,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    ms_subtitle_def_typob,
    text_def_to,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    desc_stat_subtite,
    desc_stat_rgd_content,
    desc_stat_content,
    ms_distribution_title,
    ms_tendances_temporelles,
    ms_variabilite_cameras,
    ms_ecart_type_intercam,
    data_viz,
    ms_repartition_camera,
    ms_corr_analysis,
    ms_anomali_detection,
    ms_seg_group,
    
    ms_interpretation_synthese_title,
    
    #summary 
    ms_summary_subtitle,
    ms_summary_text,
    ms_summary_result, 
            
    #interpretation
    ms_interpretation_subtitle,
    ms_interpretation_text,
    ms_interpretation_result,
    
    # lien avec l'objet du rapport
    ms_lien_rapport_subtitle,
    ms_lien_rapport,
            
    #conclusion preliminaire
    ms_conclusison_pre_subtitle,
    ms_conclusion_pre,
    
    #conclusion general
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
    ms_point_final,
):
    
    pdf_data = export_to_docx_pdf(
        title,
        subtitle,
        intro_title,
        intro_brief,
        title_objectif,
        intro_objectif,
        title_scope,
        text_scope,
        title_data_desc,
        sub_title_data_text,
        data_overview_text,
        sub_title_data_desc_col,
        ms_subtitle_def_typob,
        text_def_to,
        data_cleaning_title,
        missing_val_title,
        missing_val_content,
        transformation_data_title,
        transformation_data_content,
        desc_stat_title,
        desc_stat_subtite,
        desc_stat_rgd_content,
        desc_stat_content,
        ms_distribution_title,
        ms_tendances_temporelles,
        ms_variabilite_cameras,
        ms_ecart_type_intercam,
        data_viz,
        ms_repartition_camera,
        ms_corr_analysis,
        ms_anomali_detection,
        ms_seg_group,
        
        ms_interpretation_synthese_title,
        #summary
        ms_summary_subtitle,
        ms_summary_text,
        ms_summary_result, 
            
        #interpretation
        ms_interpretation_subtitle,
        ms_interpretation_text,
        ms_interpretation_result,
        
         # lien avec l'objet du rapport
         ms_lien_rapport_subtitle,
         ms_lien_rapport,
            
        #conclusion preliminaire
        ms_conclusison_pre_subtitle,
        ms_conclusion_pre,
        
        #conclusion
        conclusion_title,
        conclusion_subtitl_r,
        resume_content,
        recomm_title,
        recomm_content,
        ms_point_final,
    )
    return pdf_data
