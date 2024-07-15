import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml import OxmlElement
from datetime import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
import plotly.express as px
from plotly.offline import plot
import plotly.io as pio
from docx2pdf import convert
from io import BytesIO
import io


def set_column_widths(table, widths):
    """Function to set column widths in a table."""
    for i, width in enumerate(widths):
        if isinstance(width, pd.DataFrame):
            # Calculate the maximum width based on the content of the DataFrame
            max_width = Inches(width.astype(str).applymap(len).max().max())
            table.cell(0, i).width = max_width
        else:
            # Assume width is a valid length unit (e.g., Inches(1))
            table.cell(0, i).width = width


def create_custom_table_style(doc, style_name, grid_color_rgb=(0, 0, 0)):
    """Function to create a custom table style."""
    style_element = OxmlElement("w:style")
    style_element.set(qn("w:type"), "table")
    style_element.set(qn("w:styleId"), style_name)

    name_element = OxmlElement("w:name")
    name_element.set(qn("w:val"), style_name)
    style_element.append(name_element)

    based_on_element = OxmlElement("w:basedOn")
    based_on_element.set(qn("w:val"), "TableGrid")
    style_element.append(based_on_element)

    paragraph_properties_element = OxmlElement("w:pPr")
    spacing_element = OxmlElement("w:spacing")
    spacing_element.set(qn("w:before"), "0")
    spacing_element.set(qn("w:after"), "0")
    paragraph_properties_element.append(spacing_element)
    style_element.append(paragraph_properties_element)

    table_properties_element = OxmlElement("w:tblPr")
    table_style_element = OxmlElement("w:tblStyle")
    table_style_element.set(qn("w:val"), style_name)
    table_properties_element.append(table_style_element)

    grid_element = OxmlElement("w:tblBorders")
    grid_element.set(qn("w:val"), "grid")
    grid_element.set(qn("w:color"), f"auto")
    grid_element.set(qn("w:sz"), "4")
    grid_element.set(qn("w:space"), "0")
    table_properties_element.append(grid_element)

    style_element.append(table_properties_element)

    color_element = OxmlElement("w:shd")
    color_element.set(
        qn("w:fill"),
        f"{grid_color_rgb[0]:02X}{grid_color_rgb[1]:02X}{grid_color_rgb[2]:02X}",
    )
    style_element.append(color_element)

    doc.styles._element.append(style_element)
    return style_name


def export_to_docx_pdf(
    title,
    subtitle,
    intro_title,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    col_data_desc_col,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
):

    # Create a new Word document
    doc = Document("report_template.docx")

    # Set font for the entire document
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            if "Normal" in style.name:
                style.font.name = "Times New Roman"

    # Create custom table styles
    table_style_name = create_custom_table_style(doc, "TableGrid4Accent5")

    # Set margins for the section
    section = doc.sections[0]
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    # Add titles and subtitles
    def add_heading(
        doc, level, text, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True
    ):
        heading = doc.add_heading(level=level)
        heading.alignment = alignment
        run = heading.add_run(text)
        run.bold = bold

    def add_paragraph(
        doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, font_size=14, bold=False
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.bold = bold

    # Adding content
    add_heading(doc, 0, title)
    add_heading(doc, 1, subtitle)
    add_heading(doc, 1, intro_title)
    add_heading(doc, 2, title_objectif)
    add_paragraph(doc, intro_objectif)
    add_heading(doc, 2, title_scope)
    add_paragraph(doc, text_scope)
    add_heading(doc, 2, title_data_desc)
    add_heading(doc, 3, sub_title_data_text)
    add_paragraph(doc, data_overview_text)
    add_heading(doc, 3, sub_title_data_desc_col)
    add_paragraph(doc, col_data_desc_col)
    add_heading(doc, 2, data_cleaning_title)
    add_heading(doc, 3, missing_val_title)
    add_paragraph(doc, missing_val_content)
    add_heading(doc, 3, transformation_data_title)
    add_paragraph(doc, transformation_data_content)
    add_heading(doc, 2, desc_stat_title)
    add_heading(doc, 2, conclusion_title, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_heading(doc, 3, conclusion_subtitl_r, WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph(doc, resume_content)
    add_heading(doc, 3, recomm_title)
    add_paragraph(doc, recomm_content)

    # Auto-save function using report title and timestamp
    timestamp = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
    auto_save_filename = f"{title.replace(' ', '_')}-{timestamp}.docx"
    pdf_filename = f"{title.replace(' ', '_')}-{timestamp}.pdf"

    # Save the document
    doc.save(auto_save_filename)

    # Convert the DOCX to PDF
    convert(auto_save_filename, pdf_filename)

    # Return doc as buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_report(
    title,
    subtitle,
    intro_title,
    title_objectif,
    intro_objectif,
    title_scope,
    text_scope,
    title_data_desc,
    sub_title_data_text,
    data_overview_text,
    sub_title_data_desc_col,
    col_data_desc_col,
    data_cleaning_title,
    missing_val_title,
    missing_val_content,
    transformation_data_title,
    transformation_data_content,
    desc_stat_title,
    conclusion_title,
    conclusion_subtitl_r,
    resume_content,
    recomm_title,
    recomm_content,
):
    pdf_data = export_to_docx_pdf(
        title,
        subtitle,
        intro_title,
        title_objectif,
        intro_objectif,
        title_scope,
        text_scope,
        title_data_desc,
        sub_title_data_text,
        data_overview_text,
        sub_title_data_desc_col,
        col_data_desc_col,
        data_cleaning_title,
        missing_val_title,
        missing_val_content,
        transformation_data_title,
        transformation_data_content,
        desc_stat_title,
        conclusion_title,
        conclusion_subtitl_r,
        resume_content,
        recomm_title,
        recomm_content,
    )
    return pdf_data
